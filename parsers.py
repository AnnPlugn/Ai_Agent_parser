import os
import ast
import re
import astunparse
from typing import List, Dict
import PyPDF2
import logging
from docx import Document
import xml.etree.ElementTree as ET  # Для XML
import markdown  # Для MD (pip install markdown)
# Для Java: используйте javalang (pip install javalang)
from langchain_community.document_loaders import TextLoader
import javalang
import csv
import pandas as pd
from openpyxl import load_workbook

# Настройка логирования
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


def extract_dependencies(file_path: str) -> Dict[str, List[str]]:
    """Извлечение зависимостей библиотек и функций из файла кода"""
    dependencies = {"libraries": [], "functions": []}
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()
            tree = ast.parse(content)

        # Извлечение импортов
        for node in ast.walk(tree):
            if isinstance(node, (ast.Import, ast.ImportFrom)):
                for name in getattr(node, 'names', []):
                    dependencies["libraries"].append(name.name)

        # Извлечение определений функций и их вызовов
        for node in ast.walk(tree):
            if isinstance(node, ast.FunctionDef):
                dependencies["functions"].append(node.name)
            elif isinstance(node, ast.Call) and isinstance(node.func, ast.Name):
                dependencies["functions"].append(node.func.id)

        # Удаление дубликатов
        dependencies["libraries"] = list(set(dependencies["libraries"]))
        dependencies["functions"] = list(set(dependencies["functions"]))

    except Exception as e:
        logger.error(f"Ошибка при анализе зависимостей {file_path}: {e}")
    return dependencies


def load_requirements(file_path: str, batch_size: int = 1000) -> str:
    """Загрузка бизнес-требований из текстовых файлов"""
    ext = os.path.splitext(file_path)[1].lower()
    try:
        if ext == '.pdf':
            return extract_text_from_pdf(file_path, batch_size)
        elif ext == '.docx':
            return extract_text_from_docx(file_path, batch_size)
        elif ext == '.txt':
            return extract_text_from_txt(file_path, batch_size)
        elif ext == '.md':
            return parse_md(file_path, batch_size)
        elif ext == '.csv':
            return extract_csv_requirements(file_path, batch_size)
        elif ext in ['.xlsx', '.xls']:
            return extract_excel_requirements(file_path, batch_size)
        else:
            loader = TextLoader(file_path, encoding='utf-8')
            docs = loader.load()
            text = "\n".join([doc.page_content for doc in docs])
            return "\n".join([text[i:i + batch_size] for i in range(0, len(text), batch_size)])
    except Exception as e:
        logger.error(f"Ошибка при загрузке требований из {file_path}: {e}")
        return f"Ошибка загрузки: {str(e)}"


def get_project_structure(directory_path: str) -> Dict[str, List[str]]:
    """Получение архитектуры проекта"""
    structure = {}
    for root, dirs, files in os.walk(directory_path):
        if '.git' in dirs:
            dirs.remove('.git')
        relative_path = os.path.relpath(root, directory_path)
        structure[relative_path] = files
    return structure


def extract_text_from_pdf(file_path: str, batch_size: int = 1000) -> str:
    """Извлечение текста из PDF файла с батчевой обработкой"""
    try:
        with open(file_path, 'rb') as file:
            reader = PyPDF2.PdfReader(file)
            text = ""
            for page_num in range(len(reader.pages)):
                page_text = reader.pages[page_num].extract_text() or ""
                # Обрабатываем текст батчами
                for i in range(0, len(page_text), batch_size):
                    text += page_text[i:i + batch_size] + "\n"
            return text
    except Exception as e:
        logger.error(f"Ошибка при извлечении текста из PDF {file_path}: {e}")
        return f"Ошибка извлечения: {str(e)}"


def extract_text_from_docx(file_path: str, batch_size: int = 1000) -> str:
    """Извлечение текста из DOCX файла с батчевой обработкой"""
    try:
        doc = Document(file_path)
        text = ""
        full_text = "\n".join([para.text for para in doc.paragraphs])
        # Обрабатываем текст батчами
        for i in range(0, len(full_text), batch_size):
            text += full_text[i:i + batch_size] + "\n"
        return text
    except Exception as e:
        logger.error(f"Ошибка при извлечении текста из DOCX {file_path}: {e}")
        return f"Ошибка извлечения: {str(e)}"


def extract_text_from_txt(file_path: str, batch_size: int = 1000) -> str:
    """Извлечение текста из TXT файла с батчевой обработкой"""
    try:
        with open(file_path, 'r', encoding='utf-8') as file:
            content = file.read()
            text = ""
            # Обрабатываем текст батчами
            for i in range(0, len(content), batch_size):
                text += content[i:i + batch_size] + "\n"
            return text
    except Exception as e:
        logger.error(f"Ошибка при извлечении текста из TXT {file_path}: {e}")
        return f"Ошибка извлечения: {str(e)}"


def extract_csv_requirements(file_path: str, batch_size: int = 1000) -> str:
    """Извлечение требований из CSV файла"""
    try:
        requirements_text = ""

        # Читаем CSV с помощью pandas для лучшей обработки
        try:
            df = pd.read_csv(file_path, encoding='utf-8')
        except UnicodeDecodeError:
            df = pd.read_csv(file_path, encoding='cp1251')

        # Ищем колонки с требованиями, описаниями, задачами
        requirement_columns = []
        for col in df.columns:
            col_lower = col.lower()
            if any(keyword in col_lower for keyword in
                   ['требован', 'описан', 'задач', 'функци', 'специф', 'requirement', 'description', 'task', 'feature',
                    'spec']):
                requirement_columns.append(col)

        if requirement_columns:
            requirements_text += f"CSV файл содержит {len(df)} записей с требованиями:\n\n"

            for idx, row in df.iterrows():
                requirements_text += f"Запись {idx + 1}:\n"
                for col in requirement_columns:
                    if pd.notna(row[col]) and str(row[col]).strip():
                        requirements_text += f"  {col}: {str(row[col])}\n"
                requirements_text += "\n"

                # Батчевая обработка
                if len(requirements_text) >= batch_size:
                    break
        else:
            # Если специальных колонок нет, берем все текстовые данные
            requirements_text += f"CSV файл с {len(df)} записями:\n"
            requirements_text += df.to_string(max_rows=20)

        return requirements_text
    except Exception as e:
        logger.error(f"Ошибка при извлечении требований из CSV {file_path}: {e}")
        return f"Ошибка CSV: {str(e)}"


def extract_excel_requirements(file_path: str, batch_size: int = 1000) -> str:
    """Извлечение требований из Excel файла"""
    try:
        requirements_text = ""

        # Читаем Excel файл
        try:
            df = pd.read_excel(file_path, engine='openpyxl')
        except Exception:
            # Пробуем другие движки
            df = pd.read_excel(file_path, engine='xlrd')

        # Получаем информацию о листах
        xl_file = pd.ExcelFile(file_path)
        sheet_names = xl_file.sheet_names

        requirements_text += f"Excel файл с листами: {', '.join(sheet_names)}\n\n"

        # Обрабатываем каждый лист
        for sheet_name in sheet_names:
            sheet_df = pd.read_excel(file_path, sheet_name=sheet_name)

            # Ищем колонки с требованиями
            requirement_columns = []
            for col in sheet_df.columns:
                if pd.notna(col):  # Проверяем, что название колонки не NaN
                    col_lower = str(col).lower()
                    if any(keyword in col_lower for keyword in
                           ['требован', 'описан', 'задач', 'функци', 'специф', 'requirement', 'description', 'task',
                            'feature', 'spec', 'агент', 'agent', 'роль', 'role']):
                        requirement_columns.append(col)

            if requirement_columns:
                requirements_text += f"\nЛист '{sheet_name}' содержит {len(sheet_df)} записей:\n"

                for idx, row in sheet_df.iterrows():
                    if idx >= 50:  # Ограничиваем количество строк
                        break

                    has_content = False
                    row_text = f"  Запись {idx + 1}:\n"

                    for col in requirement_columns:
                        if pd.notna(row[col]) and str(row[col]).strip() and str(row[col]) != 'nan':
                            row_text += f"    {col}: {str(row[col])}\n"
                            has_content = True

                    if has_content:
                        requirements_text += row_text

                    # Батчевая обработка
                    if len(requirements_text) >= batch_size:
                        break
            else:
                # Если специальных колонок нет, берем общий обзор
                requirements_text += f"\nЛист '{sheet_name}' - общая структура:\n"
                requirements_text += str(sheet_df.head(10).to_string()) + "\n"

        return requirements_text
    except Exception as e:
        logger.error(f"Ошибка при извлечении требований из Excel {file_path}: {e}")
        return f"Ошибка Excel: {str(e)}"


def extract_prompts(file_path: str) -> List[str]:
    """
    Извлечение промптов из файлов кода и документов с обработкой сложных структур.
    Расширенная версия для поддержки различных типов файлов.
    """
    prompts = set()
    ext = os.path.splitext(file_path)[1].lower()

    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()

        # Обработка по типу файла
        if ext == '.py':
            prompts.update(extract_prompts_from_python(content, file_path))
        elif ext in ['.js', '.ts']:
            prompts.update(extract_prompts_from_javascript(content))
        elif ext in ['.java']:
            prompts.update(extract_prompts_from_java(content))
        elif ext in ['.md', '.txt']:
            prompts.update(extract_prompts_from_text(content))
        elif ext in ['.docx']:
            prompts.update(extract_prompts_from_docx(file_path))
        elif ext == '.pdf':
            prompts.update(extract_prompts_from_pdf(file_path))
        elif ext in ['.csv', '.xlsx', '.xls']:
            prompts.update(extract_prompts_from_spreadsheet(file_path))

        # Общий поиск по регулярным выражениям для всех типов файлов
        prompts.update(extract_prompts_with_regex(content))

    except Exception as e:
        logger.error(f"Ошибка при извлечении промптов из {file_path}: {e}")

    return list(prompts)


def extract_prompts_from_python(content: str, file_path: str) -> set:
    """Извлечение промптов из Python файлов"""
    prompts = set()

    try:
        tree = ast.parse(content, filename=file_path)
        for node in ast.walk(tree):
            if isinstance(node, ast.Str):
                prompt_text = astunparse.unparse(node).strip().strip('\'"')
                if is_likely_prompt(prompt_text):
                    cleaned_prompt = ' '.join(prompt_text.strip().split())
                    if len(cleaned_prompt) > 20:  # Минимальная длина промпта
                        prompts.add(cleaned_prompt)
    except SyntaxError:
        logger.warning(f"Синтаксическая ошибка в Python файле {file_path}, используется regex")

    return prompts


def extract_prompts_from_javascript(content: str) -> set:
    """Извлечение промптов из JavaScript/TypeScript файлов"""
    prompts = set()

    # Паттерны для JavaScript
    js_patterns = [
        r'prompt\s*[:=]\s*[`\'\"](.*?)[`\'\"]',
        r'template\s*[:=]\s*[`\'\"](.*?)[`\'\"]',
        r'message\s*[:=]\s*[`\'\"](.*?)[`\'\"]',
        r'system\s*[:=]\s*[`\'\"](.*?)[`\'\"]',
        r'user\s*[:=]\s*[`\'\"](.*?)[`\'\"]',
        r'assistant\s*[:=]\s*[`\'\"](.*?)[`\'\"]',
    ]

    for pattern in js_patterns:
        matches = re.findall(pattern, content, re.IGNORECASE | re.DOTALL)
        for match in matches:
            if is_likely_prompt(match):
                cleaned_prompt = ' '.join(match.strip().split())
                if len(cleaned_prompt) > 20:
                    prompts.add(cleaned_prompt)

    return prompts


def extract_prompts_from_java(content: str) -> set:
    """Извлечение промптов из Java файлов"""
    prompts = set()

    # Паттерны для Java строк
    java_patterns = [
        r'String\s+\w*[Pp]rompt\w*\s*=\s*"(.*?)";',
        r'String\s+\w*[Tt]emplate\w*\s*=\s*"(.*?)";',
        r'String\s+\w*[Mm]essage\w*\s*=\s*"(.*?)";',
        r'"(.*?(?:analyze|задача|ты|you are|assistant|system).*?)"',
    ]

    for pattern in java_patterns:
        matches = re.findall(pattern, content, re.IGNORECASE | re.DOTALL)
        for match in matches:
            if is_likely_prompt(match):
                cleaned_prompt = ' '.join(match.strip().split())
                if len(cleaned_prompt) > 20:
                    prompts.add(cleaned_prompt)

    return prompts


def extract_prompts_from_text(content: str) -> set:
    """Извлечение промптов из текстовых файлов"""
    prompts = set()

    lines = content.split('\n')
    current_prompt = ""

    for line in lines:
        line = line.strip()

        # Ищем начало промпта
        if any(keyword in line.lower() for keyword in
               ['prompt:', 'промпт:', 'system:', 'user:', 'assistant:', 'задача:', 'task:']):
            if current_prompt and is_likely_prompt(current_prompt):
                cleaned = ' '.join(current_prompt.strip().split())
                if len(cleaned) > 20:
                    prompts.add(cleaned)
            current_prompt = line
        elif current_prompt and line:
            current_prompt += " " + line
        elif current_prompt and not line:
            # Конец промпта
            if is_likely_prompt(current_prompt):
                cleaned = ' '.join(current_prompt.strip().split())
                if len(cleaned) > 20:
                    prompts.add(cleaned)
            current_prompt = ""

    # Обработка последнего промпта
    if current_prompt and is_likely_prompt(current_prompt):
        cleaned = ' '.join(current_prompt.strip().split())
        if len(cleaned) > 20:
            prompts.add(cleaned)

    return prompts


def extract_prompts_from_docx(file_path: str) -> set:
    """Извлечение промптов из DOCX файлов"""
    prompts = set()

    try:
        doc = Document(file_path)
        full_text = "\n".join([para.text for para in doc.paragraphs])
        prompts.update(extract_prompts_from_text(full_text))
    except Exception as e:
        logger.error(f"Ошибка при извлечении промптов из DOCX {file_path}: {e}")

    return prompts


def extract_prompts_from_pdf(file_path: str) -> set:
    """Извлечение промптов из PDF файлов"""
    prompts = set()

    try:
        with open(file_path, 'rb') as file:
            reader = PyPDF2.PdfReader(file)
            full_text = ""
            for page in reader.pages:
                full_text += page.extract_text() or ""

        prompts.update(extract_prompts_from_text(full_text))
    except Exception as e:
        logger.error(f"Ошибка при извлечении промптов из PDF {file_path}: {e}")

    return prompts


def extract_prompts_from_spreadsheet(file_path: str) -> set:
    """Извлечение промптов из CSV/Excel файлов"""
    prompts = set()
    ext = os.path.splitext(file_path)[1].lower()

    try:
        if ext == '.csv':
            df = pd.read_csv(file_path)
        else:
            df = pd.read_excel(file_path)

        # Ищем колонки с промптами
        prompt_columns = []
        for col in df.columns:
            if pd.notna(col):
                col_lower = str(col).lower()
                if any(keyword in col_lower for keyword in
                       ['prompt', 'промпт', 'template', 'шаблон', 'message', 'сообщение', 'system', 'user',
                        'assistant']):
                    prompt_columns.append(col)

        # Извлекаем промпты из найденных колонок
        for col in prompt_columns:
            for value in df[col].dropna():
                text = str(value)
                if is_likely_prompt(text) and len(text) > 20:
                    cleaned = ' '.join(text.strip().split())
                    prompts.add(cleaned)

        # Если специальных колонок нет, ищем во всех текстовых колонках
        if not prompt_columns:
            for col in df.columns:
                if df[col].dtype == 'object':  # Текстовые колонки
                    for value in df[col].dropna():
                        text = str(value)
                        if is_likely_prompt(text) and len(text) > 50:
                            cleaned = ' '.join(text.strip().split())
                            prompts.add(cleaned)

    except Exception as e:
        logger.error(f"Ошибка при извлечении промптов из {file_path}: {e}")

    return prompts


def extract_prompts_with_regex(content: str) -> set:
    """Общий поиск промптов с помощью регулярных выражений"""
    prompts = set()

    # Расширенные паттерны для поиска промптов
    patterns = [
        # ChatPromptTemplate и PromptTemplate
        r'ChatPromptTemplate\.from_template\s*\(\s*["\']([^"\']+)["\']',
        r'PromptTemplate\s*\([^)]*?template\s*=\s*["\']([^"\']+)["\']',

        # Переменные с prompt в названии
        r'(?:^|\s)(\w*[Pp]rompt\w*)\s*=\s*["\']([^"\']+)["\']',
        r'(?:^|\s)(\w*[Tt]emplate\w*)\s*=\s*["\']([^"\']+)["\']',

        # Многострочные строки с ключевыми словами
        r'["\'{3}](.*?(?:analyze|анализ|задача|ты|you are|assistant|system|user|instruction).*?)["\'{3}]',

        # Словари промптов
        r'PROMPTS?\s*\[\s*["\']([^"\']+)["\']\s*\]\s*=\s*["\']([^"\']+)["\']',

        # f-строки с промптами
        r'f["\']([^"\']*(?:analyze|анализ|задача|ты|you are|assistant).*?)["\']',
    ]

    for pattern in patterns:
        matches = re.findall(pattern, content, re.IGNORECASE | re.DOTALL | re.MULTILINE)
        for match in matches:
            # match может быть строкой или кортежем
            text = match if isinstance(match, str) else ' '.join(match)
            if is_likely_prompt(text):
                cleaned_prompt = ' '.join(text.strip().split())
                if len(cleaned_prompt) > 20:
                    prompts.add(cleaned_prompt)

    return prompts


def is_likely_prompt(text: str) -> bool:
    """Проверка, является ли текст вероятно промптом"""
    if len(text) < 10:
        return False

    text_lower = text.lower()

    # Ключевые слова, указывающие на промпт
    prompt_keywords = [
        'analyze', 'анализ', 'анализируй', 'проанализируй',
        'задача', 'task', 'цель', 'goal',
        'ты', 'you are', 'вы являетесь',
        'assistant', 'ассистент', 'помощник',
        'system', 'системный',
        'user', 'пользователь',
        'instruction', 'инструкция',
        'prompt', 'промпт',
        'template', 'шаблон',
        'генерируй', 'generate',
        'создай', 'create',
        'опиши', 'describe',
        'объясни', 'explain',
        'сформируй', 'form'
    ]

    # Исключения (не промпты)
    exclude_keywords = [
        'import ', 'from ', 'class ', 'def ', 'return',
        'print(', 'console.log', 'logger.',
        'http://', 'https://', 'www.',
        '#!/', '#include', '/*', '*/',
    ]

    # Проверка на исключения
    for exclude in exclude_keywords:
        if exclude in text_lower:
            return False

    # Проверка на ключевые слова промптов
    for keyword in prompt_keywords:
        if keyword in text_lower:
            return True

    # Дополнительная проверка: если текст содержит {переменные} или инструкции
    if ('{' in text and '}' in text) or any(word in text_lower for word in ['следующ', 'based on', 'на основе']):
        return True

    return False


# Существующие функции (без изменений)
def extract_java_dependencies(file_path: str) -> Dict[str, List[str]]:
    """Парсер для Java файлов с использованием javalang"""
    dependencies = {"libraries": [], "functions": []}
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()

        # Парсим Java код
        tree = javalang.parse.parse(content)

        # Извлекаем импорты
        for import_decl in tree.imports:
            dependencies["libraries"].append(import_decl.path)

        # Извлекаем методы (функции)
        for path, node in tree:
            if isinstance(node, javalang.tree.MethodDeclaration):
                dependencies["functions"].append(node.name)
            # Также можно добавить вызовы методов, если нужно
            elif isinstance(node, javalang.tree.MethodInvocation):
                if hasattr(node, 'member') and node.member:
                    dependencies["functions"].append(node.member)

        # Удаляем дубликаты
        dependencies["libraries"] = list(set(dependencies["libraries"]))
        dependencies["functions"] = list(set(dependencies["functions"]))

    except javalang.parser.JavaSyntaxError as e:
        logger.warning(f"Синтаксическая ошибка в Java файле {file_path}: {e}")
    except Exception as e:
        logger.error(f"Ошибка при анализе Java зависимостей {file_path}: {e}")

    return dependencies


def extract_js_dependencies(file_path: str) -> Dict[str, List[str]]:
    """Regex-парсер для JS (библиотеки: import/from, функции: function/def)"""
    dependencies = {"libraries": [], "functions": []}
    with open(file_path, 'r', encoding='utf-8') as f:
        content = f.read()
    # Импорты: import ... from 'lib';
    lib_matches = re.findall(r"import\s+.*?\s+from\s+['\"](.*?)['\"]", content)
    dependencies["libraries"] = list(set(lib_matches))
    # Функции: function name() или const name = ()
    func_matches = re.findall(r"(?:function|const|let|var)\s+([a-zA-Z_]\w*)\s*(?:=|\()", content)
    dependencies["functions"] = list(set(func_matches))
    return dependencies


def extract_cpp_dependencies(file_path: str) -> Dict[str, List[str]]:
    """Regex для C++: #include <lib>, функции"""
    dependencies = {"libraries": [], "functions": []}
    with open(file_path, 'r', encoding='utf-8') as f:
        content = f.read()
    lib_matches = re.findall(r"#include\s+[<\"](.*?)[\">]", content)
    dependencies["libraries"] = list(set(lib_matches))
    func_matches = re.findall(r"(\w+)\s*\([^)]*\)\s*\{", content)  # Простой паттерн для функций
    dependencies["functions"] = list(set(func_matches))
    return dependencies


def smart_parse_txt_bpmn(file_path: str, batch_size: int = 1000) -> str:
    """Умный парсинг для TXT и BPMN (low-code). Если XML внутри – парсить как XML."""
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()
        if '<' in content and '>' in content and '<?xml' in content:  # Проверка на XML
            # Low-code: парсим как XML
            root = ET.fromstring(content)
            text = ET.tostring(root, encoding='unicode', method='text')  # Извлекаем текст
            # Или более умно: извлекаем атрибуты/элементы для BPMN (процессы, задачи)
            processes = [elem.attrib.get('name', '') for elem in root.findall('.//{*}process')]
            tasks = [elem.attrib.get('name', '') for elem in root.findall('.//{*}task')]
            text += f"\nBPMN Processes: {', '.join(processes)}\nTasks: {', '.join(tasks)}"
        else:
            text = content  # Обычный TXT
        # Батчевая обработка
        batched = ""
        for i in range(0, len(text), batch_size):
            batched += text[i:i + batch_size] + "\n"
        return batched
    except Exception as e:
        logging.error(f"Ошибка парсинга TXT/BPMN {file_path}: {e}")
        return f"Ошибка: {str(e)}"


def parse_md(file_path: str, batch_size: int = 1000) -> str:
    """Отдельный парсер для MD: конверт в HTML или извлекай структуру."""
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            md_content = f.read()
        # Используем markdown lib для конвертации в HTML (или извлечения)
        html = markdown.markdown(md_content)
        # Или структура: извлекаем заголовки
        headers = re.findall(r'^#{1,6}\s+(.*)', md_content, re.MULTILINE)
        text = f"MD Headers: {', '.join(headers)}\n\n{html}"
        # Батчевая
        batched = ""
        for i in range(0, len(text), batch_size):
            batched += text[i:i + batch_size] + "\n"
        return batched
    except Exception as e:
        logging.error(f"Ошибка парсинга MD {file_path}: {e}")
        return f"Ошибка: {str(e)}"


def extract_business_requirements(file_path: str) -> str:
    """Извлечение бизнес-требований из различных типов файлов"""
    ext = os.path.splitext(file_path)[1].lower()

    try:
        if ext in ['.docx']:
            return extract_business_from_docx(file_path)
        elif ext == '.pdf':
            return extract_business_from_pdf(file_path)
        elif ext in ['.txt', '.md']:
            return extract_business_from_text(file_path)
        elif ext == '.csv':
            return extract_business_from_csv(file_path)
        elif ext in ['.xlsx', '.xls']:
            return extract_business_from_excel(file_path)
        else:
            return load_requirements(file_path)
    except Exception as e:
        logger.error(f"Ошибка извлечения бизнес-требований из {file_path}: {e}")
        return f"Ошибка: {str(e)}"


def extract_business_from_docx(file_path: str) -> str:
    """Извлечение бизнес-требований из DOCX"""
    try:
        doc = Document(file_path)
        business_content = ""

        # Ищем разделы с бизнес-требованиями
        business_keywords = [
            'требования', 'requirements', 'бизнес', 'business',
            'функциональные', 'functional', 'нефункциональные', 'non-functional',
            'пользовательские истории', 'user stories', 'use case',
            'процесс', 'process', 'workflow', 'сценарий', 'scenario',
            'описание системы', 'system description', 'архитектура', 'architecture'
        ]

        current_section = ""
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            # Проверяем, является ли это заголовком секции
            if any(keyword in text.lower() for keyword in business_keywords):
                if current_section:
                    business_content += f"\n\n--- {current_section} ---\n"
                current_section = text

            # Добавляем содержимое
            if current_section:
                business_content += text + "\n"

        return business_content if business_content else extract_text_from_docx(file_path)
    except Exception as e:
        logger.error(f"Ошибка извлечения бизнес-требований из DOCX {file_path}: {e}")
        return f"Ошибка: {str(e)}"


def extract_business_from_pdf(file_path: str) -> str:
    """Извлечение бизнес-требований из PDF"""
    try:
        full_text = extract_text_from_pdf(file_path)

        # Ищем разделы с бизнес-требованиями
        business_sections = []
        lines = full_text.split('\n')

        current_section = ""
        is_business_section = False

        business_keywords = [
            'требования', 'requirements', 'бизнес', 'business',
            'функциональные', 'functional', 'техническое задание', 'tz',
            'спецификация', 'specification', 'архитектура', 'architecture'
        ]

        for line in lines:
            line = line.strip()
            if not line:
                continue

            # Проверяем начало бизнес-секции
            if any(keyword in line.lower() for keyword in business_keywords):
                if current_section and is_business_section:
                    business_sections.append(current_section)
                current_section = line + "\n"
                is_business_section = True
            elif is_business_section:
                current_section += line + "\n"
                # Если встречаем новый заголовок, завершаем текущую секцию
                if line.isupper() or (len(line) < 50 and line.endswith(':')):
                    is_business_section = False

        if current_section and is_business_section:
            business_sections.append(current_section)

        return "\n\n".join(business_sections) if business_sections else full_text
    except Exception as e:
        logger.error(f"Ошибка извлечения бизнес-требований из PDF {file_path}: {e}")
        return f"Ошибка: {str(e)}"


def extract_business_from_text(file_path: str) -> str:
    """Извлечение бизнес-требований из текстовых файлов"""
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()

        # Для MD файлов используем специальную обработку
        if file_path.endswith('.md'):
            # Ищем заголовки с бизнес-содержимым
            sections = []
            current_section = ""

            for line in content.split('\n'):
                if line.startswith('#'):
                    if current_section:
                        sections.append(current_section)
                    current_section = line + "\n"
                else:
                    current_section += line + "\n"

            if current_section:
                sections.append(current_section)

            # Фильтруем секции с бизнес-содержимым
            business_sections = []
            business_keywords = ['требования', 'business', 'архитектура', 'описание', 'функции', 'features']

            for section in sections:
                if any(keyword in section.lower() for keyword in business_keywords):
                    business_sections.append(section)

            return "\n\n".join(business_sections) if business_sections else content

        return content
    except Exception as e:
        logger.error(f"Ошибка извлечения бизнес-требований из {file_path}: {e}")
        return f"Ошибка: {str(e)}"


def extract_business_from_csv(file_path: str) -> str:
    """Извлечение бизнес-требований из CSV"""
    try:
        df = pd.read_csv(file_path, encoding='utf-8')

        business_content = f"CSV файл с бизнес-данными ({len(df)} записей):\n\n"

        # Ищем колонки с бизнес-информацией
        business_columns = []
        for col in df.columns:
            col_lower = col.lower()
            if any(keyword in col_lower for keyword in [
                'требование', 'requirement', 'описание', 'description',
                'функция', 'function', 'процесс', 'process', 'роль', 'role',
                'задача', 'task', 'цель', 'goal', 'бизнес', 'business'
            ]):
                business_columns.append(col)

        if business_columns:
            for idx, row in df.iterrows():
                business_content += f"Запись {idx + 1}:\n"
                for col in business_columns:
                    if pd.notna(row[col]):
                        business_content += f"  {col}: {row[col]}\n"
                business_content += "\n"

                if idx >= 100:  # Ограничиваем количество записей
                    break
        else:
            # Если специальных колонок нет, показываем общую структуру
            business_content += "Структура данных:\n"
            business_content += df.head(10).to_string()

        return business_content
    except Exception as e:
        logger.error(f"Ошибка извлечения бизнес-требований из CSV {file_path}: {e}")
        return f"Ошибка: {str(e)}"


def extract_business_from_excel(file_path: str) -> str:
    """Извлечение бизнес-требований из Excel"""
    try:
        business_content = ""
        xl_file = pd.ExcelFile(file_path)

        for sheet_name in xl_file.sheet_names:
            df = pd.read_excel(file_path, sheet_name=sheet_name)

            business_content += f"\n=== Лист: {sheet_name} ===\n"

            # Ищем колонки с бизнес-информацией
            business_columns = []
            for col in df.columns:
                if pd.notna(col):
                    col_lower = str(col).lower()
                    if any(keyword in col_lower for keyword in [
                        'требование', 'requirement', 'описание', 'description',
                        'функция', 'function', 'процесс', 'process', 'роль', 'role',
                        'задача', 'task', 'цель', 'goal', 'бизнес', 'business',
                        'агент', 'agent', 'сценарий', 'scenario'
                    ]):
                        business_columns.append(col)

            if business_columns:
                business_content += f"Найдены бизнес-колонки: {', '.join(business_columns)}\n\n"

                for idx, row in df.iterrows():
                    if idx >= 50:  # Ограничиваем количество записей
                        break

                    has_content = False
                    row_content = f"Запись {idx + 1}:\n"

                    for col in business_columns:
                        if pd.notna(row[col]) and str(row[col]).strip():
                            row_content += f"  {col}: {row[col]}\n"
                            has_content = True

                    if has_content:
                        business_content += row_content + "\n"
            else:
                business_content += "Общая структура листа:\n"
                business_content += df.head(5).to_string() + "\n"

        return business_content
    except Exception as e:
        logger.error(f"Ошибка извлечения бизнес-требований из Excel {file_path}: {e}")
        return f"Ошибка: {str(e)}"


def extract_project_description(file_path: str) -> str:
    """Извлечение описания проекта из различных файлов"""
    ext = os.path.splitext(file_path)[1].lower()
    filename = os.path.basename(file_path).lower()

    # Определяем, содержит ли файл описание проекта
    description_indicators = [
        'readme', 'description', 'описание', 'about',
        'project', 'проект', 'overview', 'обзор',
        'spec', 'specification', 'техзадание', 'tz'
    ]

    if not any(indicator in filename for indicator in description_indicators):
        return ""

    try:
        if ext == '.md':
            return extract_description_from_md(file_path)
        elif ext == '.txt':
            return extract_description_from_txt(file_path)
        elif ext == '.docx':
            return extract_description_from_docx(file_path)
        elif ext == '.pdf':
            return extract_description_from_pdf(file_path)
        else:
            return load_requirements(file_path)
    except Exception as e:
        logger.error(f"Ошибка извлечения описания проекта из {file_path}: {e}")
        return f"Ошибка: {str(e)}"


def extract_description_from_md(file_path: str) -> str:
    """Извлечение описания из Markdown файла"""
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()

        # Для README файлов извлекаем основные секции
        sections = {}
        current_section = "overview"
        current_content = ""

        for line in content.split('\n'):
            if line.startswith('#'):
                if current_content.strip():
                    sections[current_section] = current_content.strip()
                current_section = line.strip('# ').lower()
                current_content = ""
            else:
                current_content += line + "\n"

        if current_content.strip():
            sections[current_section] = current_content.strip()

        # Приоритезируем важные секции
        priority_sections = ['overview', 'описание', 'description', 'about', 'введение', 'installation', 'usage']
        result = ""

        for section_key in priority_sections:
            for key, content in sections.items():
                if section_key in key:
                    result += f"\n=== {key.title()} ===\n{content}\n"

        # Добавляем остальные секции
        for key, content in sections.items():
            if not any(priority in key for priority in priority_sections):
                result += f"\n=== {key.title()} ===\n{content}\n"

        return result if result else content
    except Exception as e:
        logger.error(f"Ошибка извлечения описания из MD {file_path}: {e}")
        return f"Ошибка: {str(e)}"


def extract_description_from_txt(file_path: str) -> str:
    """Извлечение описания из TXT файла"""
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()

        # Ищем структурированное описание
        lines = content.split('\n')
        description_parts = []
        current_part = ""

        for line in lines:
            line = line.strip()
            if not line:
                if current_part:
                    description_parts.append(current_part)
                    current_part = ""
                continue

            # Проверяем, является ли строка заголовком
            if (line.isupper() or line.endswith(':') or
                    any(keyword in line.lower() for keyword in ['описание', 'функции', 'возможности', 'features'])):
                if current_part:
                    description_parts.append(current_part)
                current_part = f"=== {line} ===\n"
            else:
                current_part += line + "\n"

        if current_part:
            description_parts.append(current_part)

        return "\n\n".join(description_parts) if description_parts else content
    except Exception as e:
        logger.error(f"Ошибка извлечения описания из TXT {file_path}: {e}")
        return f"Ошибка: {str(e)}"


def extract_description_from_docx(file_path: str) -> str:
    """Извлечение описания из DOCX файла"""
    try:
        doc = Document(file_path)
        description_content = ""

        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                # Проверяем стиль параграфа для выделения заголовков
                if para.style.name.startswith('Heading'):
                    description_content += f"\n=== {text} ===\n"
                else:
                    description_content += text + "\n"

        return description_content
    except Exception as e:
        logger.error(f"Ошибка извлечения описания из DOCX {file_path}: {e}")
        return f"Ошибка: {str(e)}"


def extract_description_from_pdf(file_path: str) -> str:
    """Извлечение описания из PDF файла"""
    try:
        return extract_text_from_pdf(file_path)
    except Exception as e:
        logger.error(f"Ошибка извлечения описания из PDF {file_path}: {e}")
        return f"Ошибка: {str(e)}"


def get_supported_extensions() -> Dict[str, List[str]]:
    """Возвращает список поддерживаемых расширений файлов по категориям"""
    return {
        "code_files": ['.py', '.js', '.ts', '.java', '.cpp', '.c', '.h', '.hpp'],
        "document_files": ['.docx', '.pdf', '.txt', '.md', '.rtf'],
        "data_files": ['.csv', '.xlsx', '.xls', '.json', '.yaml', '.yml'],
        "config_files": ['.ini', '.conf', '.config', '.env'],
        "markup_files": ['.xml', '.html', '.htm', '.bpmn']
    }


def should_process_file(file_path: str) -> bool:
    """Проверяет, должен ли файл быть обработан"""
    ext = os.path.splitext(file_path)[1].lower()
    supported = get_supported_extensions()

    all_supported = []
    for category in supported.values():
        all_supported.extend(category)

    return ext in all_supported